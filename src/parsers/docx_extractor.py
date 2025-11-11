"""
DOCX text extraction utilities.

This module provides functionality to extract text from .docx files
with error handling and structure preservation.
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError(
        "python-docx is not installed. Install it with: pip install python-docx"
    )

# Set up logging
logger = logging.getLogger(__name__)


class DocxExtractor:
    """
    Extract text content from .docx files.

    Features:
        - Extracts paragraphs and tables
        - Preserves document structure
        - Handles corrupted files gracefully
        - Validates file format
    """

    def __init__(self, preserve_formatting: bool = False):
        """
        Initialize DOCX extractor.

        Args:
            preserve_formatting: If True, preserves bullets and structure
        """
        self.preserve_formatting = preserve_formatting

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a .docx file.

        Args:
            file_path: Path to the .docx file

        Returns:
            str: Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid .docx
        """
        # Validate file
        validation_error = self.validate_file(file_path)
        if validation_error:
            raise ValueError(validation_error)

        try:
            document = Document(file_path)
            text_parts = []

            # Extract text from document
            for element in document.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, document)
                    text = self._extract_paragraph_text(paragraph)
                    if text.strip():
                        text_parts.append(text)

                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, document)
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        text_parts.append(table_text)

            extracted_text = "\n".join(text_parts)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from {file_path}")

            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from document: {str(e)}")

    def extract_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text along with document metadata.

        Args:
            file_path: Path to the .docx file

        Returns:
            Dict containing:
                - text: Extracted text
                - metadata: Document metadata (author, created, etc.)
                - stats: Statistics (word count, paragraph count, etc.)

        Useful for:
            - Debugging parsing issues
            - Quality assessment
            - Audit trails
        """
        validation_error = self.validate_file(file_path)
        if validation_error:
            return {
                "text": "",
                "metadata": {},
                "stats": {},
                "error": validation_error,
            }

        try:
            document = Document(file_path)
            text = self.extract_text(file_path)

            # Extract metadata
            metadata = self._extract_metadata(document)

            # Calculate statistics
            stats = self._calculate_stats(document, text)

            return {
                "text": text,
                "metadata": metadata,
                "stats": stats,
                "error": None,
            }

        except Exception as e:
            logger.error(f"Error extracting with metadata from {file_path}: {str(e)}")
            return {
                "text": "",
                "metadata": {},
                "stats": {},
                "error": str(e),
            }

    def validate_file(self, file_path: str) -> Optional[str]:
        """
        Validate that the file exists and is a valid .docx.

        Args:
            file_path: Path to the file

        Returns:
            Optional[str]: Error message if invalid, None if valid
        """
        path = Path(file_path)

        # Check existence
        if not path.exists():
            return f"File not found: {file_path}"

        # Check if it's a file
        if not path.is_file():
            return f"Path is not a file: {file_path}"

        # Check extension
        if path.suffix.lower() not in [".docx", ".doc"]:
            return f"File is not a .docx or .doc file: {file_path}"

        # Check file size (warn if > 10MB)
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > 10:
            logger.warning(f"Large file detected: {file_size_mb:.2f}MB")

        # Check if it's a valid docx (try to open)
        try:
            Document(file_path)
        except Exception as e:
            return f"Invalid or corrupted .docx file: {str(e)}"

        return None

    # ==========================================
    # Private Helper Methods
    # ==========================================

    def _extract_paragraph_text(self, paragraph: Paragraph) -> str:
        """Extract text from a paragraph, preserving formatting if enabled"""
        text = paragraph.text

        if self.preserve_formatting and text.strip():
            # Check if it's a bullet point
            if self._is_bullet_point(paragraph):
                text = f"• {text}"

        return text

    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a table"""
        table_texts = []

        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)

            if row_texts:
                table_texts.append(" | ".join(row_texts))

        return "\n".join(table_texts)

    def _is_bullet_point(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a bullet point"""
        # Check paragraph style
        if paragraph.style and "List" in paragraph.style.name:
            return True

        # Check for bullet characters
        text = paragraph.text.strip()
        if text and text[0] in ["•", "·", "-", "◦", "▪", "▫"]:
            return True

        return False

    def _extract_metadata(self, document: Document) -> Dict[str, Any]:
        """Extract document metadata"""
        try:
            core_properties = document.core_properties

            return {
                "author": core_properties.author or "Unknown",
                "created": str(core_properties.created) if core_properties.created else None,
                "modified": str(core_properties.modified) if core_properties.modified else None,
                "title": core_properties.title or "",
                "subject": core_properties.subject or "",
            }
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
            return {}

    def _calculate_stats(self, document: Document, text: str) -> Dict[str, int]:
        """Calculate document statistics"""
        try:
            return {
                "num_paragraphs": len(document.paragraphs),
                "num_tables": len(document.tables),
                "num_characters": len(text),
                "num_words": len(text.split()),
                "num_lines": len(text.split("\n")),
            }
        except Exception as e:
            logger.warning(f"Could not calculate stats: {str(e)}")
            return {}


# ==========================================
# Batch Processing
# ==========================================

def extract_text_from_files(file_paths: List[str]) -> Dict[str, str]:
    """
    Extract text from multiple .docx files.

    Args:
        file_paths: List of file paths

    Returns:
        Dict mapping file path to extracted text

    Example:
        texts = extract_text_from_files(["resume1.docx", "resume2.docx"])
        for path, text in texts.items():
            print(f"{path}: {len(text)} characters")
    """
    extractor = DocxExtractor()
    results = {}

    for file_path in file_paths:
        try:
            text = extractor.extract_text(file_path)
            results[file_path] = text
        except Exception as e:
            logger.error(f"Failed to extract from {file_path}: {str(e)}")
            results[file_path] = ""

    return results


# ==========================================
# Main for Testing
# ==========================================

if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <path_to_docx_file>")
        sys.exit(1)

    file_path = sys.argv[1]

    extractor = DocxExtractor(preserve_formatting=True)

    # Extract with metadata
    result = extractor.extract_with_metadata(file_path)

    if result["error"]:
        print(f"Error: {result['error']}")
    else:
        print("=" * 50)
        print("METADATA:")
        print("=" * 50)
        for key, value in result["metadata"].items():
            print(f"{key}: {value}")

        print("\n" + "=" * 50)
        print("STATISTICS:")
        print("=" * 50)
        for key, value in result["stats"].items():
            print(f"{key}: {value}")

        print("\n" + "=" * 50)
        print("EXTRACTED TEXT:")
        print("=" * 50)
        print(result["text"][:500] + "..." if len(result["text"]) > 500 else result["text"])
